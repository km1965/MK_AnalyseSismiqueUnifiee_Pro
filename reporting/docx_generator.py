"""Génération de la note de calcul au format DOCX (python-docx) à partir d'un Report.

Inclut les tableaux par section et un croquis de coupe (si fourni).
"""
from __future__ import annotations

import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.results import Report

logger = logging.getLogger(__name__)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)


def _style_verdict(cell_text, verdict):
    return verdict


def generer_note_calcul_docx(nom_fichier: str, rapport: Report, sketch_path: str = None):
    doc = Document()

    doc.add_heading("Note de Calcul d'Analyse Sismique", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Titre : {rapport.title}").italic = True
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    if sketch_path:
        try:
            doc.add_picture(sketch_path, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph("Croquis de coupe (pré-dimensionnement)")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = GREY
        except Exception as e:  # noqa: BLE001
            logger.warning("Croquis non inclus : %s", e)

    global_ok = True
    for title, items in rapport.sections:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Grandeur"
        hdr[1].text = "Valeur"
        hdr[2].text = "Unité"
        hdr[3].text = "Vérification / Note"
        for it in items:
            row = table.add_row().cells
            row[0].text = it.label
            row[1].text = it.formatted()
            row[2].text = it.unit
            if it.verdict:
                r = row[3].paragraphs[0].add_run(it.verdict)
                r.bold = True
                r.font.color.rgb = GREEN if it.verdict == "OK" else RED
                if it.verdict == "NON OK":
                    global_ok = False
            else:
                row[3].text = it.note

    concl = doc.add_paragraph()
    run = concl.add_run("Conclusion : " + ("TOUTES LES VÉRIFICATIONS SONT CONFORMES."
                                            if global_ok else
                                            "DES VÉRIFICATIONS NE SONT PAS CONFORMES – revoir le dimensionnement."))
    run.bold = True
    run.font.color.rgb = GREEN if global_ok else RED

    doc.save(nom_fichier)
    logger.info("Note de calcul DOCX générée : %s", nom_fichier)
